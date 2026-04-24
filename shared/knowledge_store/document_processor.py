"""文档处理管线 — 解析 + 分块 + 向量化 + 写入

支持文件类型：
- PDF（使用 pdfplumber 或 unstructured）
- DOCX（使用 python-docx）
- XLSX（使用 openpyxl）
- TXT/MD（直接读取）
- 手动输入文本
"""

from __future__ import annotations

import hashlib
from pathlib import Path
from typing import Any

import structlog

from shared.vector_store.embeddings import EmbeddingService

from .chunker import semantic_chunk
from .pg_vector_store import PgVectorStore

logger = structlog.get_logger()


class DocumentProcessor:
    """文档处理管线"""

    @staticmethod
    async def process_document(
        document_id: str,
        file_path: str | None,
        source_type: str,
        collection: str,
        tenant_id: str,
        db: Any,
        raw_text: str | None = None,
        metadata: dict[str, Any] | None = None,
        chunk_size: int = 512,
        chunk_overlap: int = 64,
    ) -> dict[str, Any]:
        """处理文档全流程：解析 → 分块 → 向量化 → 写入。

        Args:
            document_id: 文档 ID
            file_path: 文件路径（source_type != 'manual' 时必填）
            source_type: 来源类型（manual/pdf/docx/xlsx/txt）
            collection: 知识集合名称
            tenant_id: 租户ID
            db: SQLAlchemy AsyncSession
            raw_text: 手动输入的文本（source_type == 'manual' 时使用）
            metadata: 额外元数据
            chunk_size: 分块大小（token 数）
            chunk_overlap: 分块重叠（token 数）

        Returns:
            {chunk_count: int, status: str, errors: list[str]}
        """
        errors: list[str] = []

        # 1. 解析文件获取文本
        if source_type == "manual" and raw_text:
            text = raw_text
        elif file_path:
            text = _parse_file(file_path, source_type)
            if not text:
                return {"chunk_count": 0, "status": "failed", "errors": ["文件解析失败或内容为空"]}
        else:
            return {"chunk_count": 0, "status": "failed", "errors": ["缺少文件路径或文本内容"]}

        # 2. 分块
        chunk_results = semantic_chunk(text, max_tokens=chunk_size, overlap_tokens=chunk_overlap)
        if not chunk_results:
            return {"chunk_count": 0, "status": "failed", "errors": ["分块结果为空"]}

        logger.info(
            "document_chunked",
            document_id=document_id,
            chunk_count=len(chunk_results),
        )

        # 3. 批量向量化
        texts = [c.text for c in chunk_results]
        embeddings = await EmbeddingService.embed_batch(texts)

        # 4. 构建写入数据
        chunks_to_write: list[dict[str, Any]] = []
        for i, (cr, emb) in enumerate(zip(chunk_results, embeddings)):
            doc_id = f"{collection}:{document_id}:{i}"
            chunks_to_write.append(
                {
                    "text": cr.text,
                    "embedding": emb,
                    "metadata": metadata or {},
                    "doc_id": doc_id,
                    "document_id": document_id,
                    "collection": collection,
                    "chunk_index": i,
                    "token_count": cr.token_count,
                }
            )

        # 5. 写入 pgvector
        result = await PgVectorStore.upsert_chunks(chunks_to_write, tenant_id, db)

        # 6. 更新文档状态
        try:
            from sqlalchemy import text as sql_text

            await db.execute(
                sql_text("""
                UPDATE knowledge_documents
                SET chunk_count = :chunk_count,
                    status = :status,
                    updated_at = NOW()
                WHERE id = :doc_id::uuid
                AND tenant_id = :tenant_id::uuid
            """),
                {
                    "chunk_count": result["success"],
                    "status": "published" if result["failed"] == 0 else "failed",
                    "doc_id": document_id,
                    "tenant_id": tenant_id,
                },
            )
            await db.commit()
        except Exception as exc:
            errors.append(f"更新文档状态失败: {str(exc)}")

        status = "published" if result["failed"] == 0 else "failed"
        if result["failed"] > 0:
            errors.append(f"{result['failed']}个块写入失败")

        return {
            "chunk_count": result["success"],
            "status": status,
            "errors": errors,
        }

    @staticmethod
    def compute_file_hash(file_path: str) -> str:
        """计算文件 SHA-256 哈希（用于去重）"""
        h = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                h.update(chunk)
        return h.hexdigest()


def _parse_file(file_path: str, source_type: str) -> str:
    """解析文件为纯文本。

    支持：pdf, docx, xlsx, txt, md
    """
    path = Path(file_path)
    if not path.exists():
        logger.warning("parse_file_not_found", file_path=file_path)
        return ""

    ext = path.suffix.lower()
    source = source_type.lower()

    try:
        # TXT / Markdown
        if ext in (".txt", ".md") or source in ("txt", "md", "text"):
            return path.read_text(encoding="utf-8")

        # PDF
        if ext == ".pdf" or source == "pdf":
            return _parse_pdf(file_path)

        # DOCX
        if ext == ".docx" or source == "docx":
            return _parse_docx(file_path)

        # XLSX / CSV
        if ext in (".xlsx", ".csv") or source in ("xlsx", "csv"):
            return _parse_xlsx(file_path)

        # 默认当纯文本
        return path.read_text(encoding="utf-8")

    except Exception as exc:
        logger.warning("parse_file_failed", file_path=file_path, error=str(exc), exc_info=True)
        return ""


def _parse_pdf(file_path: str) -> str:
    """解析 PDF 文件"""
    try:
        import pdfplumber

        texts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.append(text)

                # 提取表格
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join(str(cell or "") for cell in row)
                        texts.append(row_text)

        return "\n\n".join(texts)
    except ImportError:
        logger.warning("pdfplumber_not_installed")
        return ""


def _parse_docx(file_path: str) -> str:
    """解析 DOCX 文件"""
    try:
        from docx import Document

        doc = Document(file_path)
        texts = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells)
                texts.append(row_text)

        return "\n\n".join(texts)
    except ImportError:
        logger.warning("python_docx_not_installed")
        return ""


def _parse_xlsx(file_path: str) -> str:
    """解析 Excel/CSV 文件"""
    ext = Path(file_path).suffix.lower()

    if ext == ".csv":
        import csv

        texts = []
        with open(file_path, encoding="utf-8") as f:
            reader = csv.reader(f)
            for row in reader:
                texts.append(" | ".join(row))
        return "\n".join(texts)

    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, read_only=True)
        texts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell or "") for cell in row)
                if row_text.strip():
                    texts.append(row_text)
        return "\n".join(texts)
    except ImportError:
        logger.warning("openpyxl_not_installed")
        return ""
